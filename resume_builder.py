import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# Load JSON resume
with open("Manith_Luthria_Resume.json", "r") as f:
    resume = json.load(f)

# Create a Word document
doc = Document()

# Set page margins for compact layout
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Name & Contact - Compact format
name_heading = doc.add_heading(resume["name"], level=0)
name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Style the name heading
for run in name_heading.runs:
    run.font.size = Pt(18)
    run.font.bold = True

# Reduce spacing after name heading
name_heading.paragraph_format.space_after = Pt(20)

# Contact info as a heading below the name
contact = resume["contact"]
contact_heading = doc.add_heading(f"{contact['email']} | {contact['phone']} | {contact['linkedin']}", level=1)
contact_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in contact_heading.runs:
    run.font.size = Pt(10)

# Summary
summary_heading = doc.add_heading("Summary", level=1)
for run in summary_heading.runs:
    run.font.size = Pt(12)
    run.font.bold = True

summary_para = doc.add_paragraph(resume["summary"])
summary_para.runs[0].font.size = Pt(10)

# Experience
exp_heading = doc.add_heading("Experience", level=1)
for run in exp_heading.runs:
    run.font.size = Pt(12)
    run.font.bold = True

for exp in resume["experience"]:
    # Combine role, company, location, and date into one line
    job_para = doc.add_paragraph()
    job_para.add_run(f"{exp['title']} ").bold = True
    job_para.add_run(f"- {exp['company']} ").italic = True
    job_para.add_run(f"| {exp['location']} | {exp['date']}")
    job_para.runs[0].font.size = Pt(11)
    job_para.runs[1].font.size = Pt(11)
    job_para.runs[2].font.size = Pt(11)
    
    # Achievements with real bullet points (●)
    for ach in exp["achievements"]:
        bullet_para = doc.add_paragraph()
        bullet_para.add_run("● ").font.size = Pt(9)
        bullet_para.add_run(ach).font.size = Pt(9)
        # Remove extra spacing between bullet points
        bullet_para.paragraph_format.space_after = Pt(0)

# Projects
proj_heading = doc.add_heading("Projects", level=1)
for run in proj_heading.runs:
    run.font.size = Pt(12)
    run.font.bold = True

for proj in resume["projects"]:
    # Project name
    proj_para = doc.add_paragraph()
    proj_para.add_run(proj["name"]).bold = True
    proj_para.runs[0].font.size = Pt(11)
    
    # Project date
    date_para = doc.add_paragraph(proj["date"])
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].italic = True
    
    # Project details with real bullet points (●)
    for detail in proj["details"]:
        bullet_para = doc.add_paragraph()
        bullet_para.add_run("● ").font.size = Pt(9)
        bullet_para.add_run(detail).font.size = Pt(9)
        # Remove extra spacing between bullet points
        bullet_para.paragraph_format.space_after = Pt(0)

# Education
edu_heading = doc.add_heading("Education", level=1)
for run in edu_heading.runs:
    run.font.size = Pt(12)
    run.font.bold = True

for edu in resume["education"]:
    # Degree
    degree_para = doc.add_paragraph()
    degree_para.add_run(edu["degree"]).bold = True
    degree_para.runs[0].font.size = Pt(11)
    
    # School, location, year, GPA
    info_para = doc.add_paragraph(f"{edu['school']} | {edu['location']} | {edu['year']} | GPA: {edu['gpa']}")
    info_para.runs[0].font.size = Pt(9)
    info_para.runs[0].italic = True
    
    # Relevant coursework inline with commas
    course_para = doc.add_paragraph()
    course_para.add_run("Relevant Coursework: ").bold = True
    course_para.add_run(", ".join(edu["relevant_coursework"]))
    course_para.runs[0].font.size = Pt(9)
    course_para.runs[1].font.size = Pt(9)

# Skills - inline format
skills_heading = doc.add_heading("Skills", level=1)
for run in skills_heading.runs:
    run.font.size = Pt(12)
    run.font.bold = True

skills_para = doc.add_paragraph()
skills_para.add_run("Languages: ").bold = True
skills_para.add_run(", ".join(resume["skills"]["languages"]))
skills_para.runs[0].font.size = Pt(10)
skills_para.runs[1].font.size = Pt(10)

tools_para = doc.add_paragraph()
tools_para.add_run("Tools & Technologies: ").bold = True
tools_para.add_run(", ".join(resume["skills"]["tools_and_technologies"]))
tools_para.runs[0].font.size = Pt(10)
tools_para.runs[1].font.size = Pt(10)

# Save Word doc
doc.save("Manith_Luthria_Resume.docx")